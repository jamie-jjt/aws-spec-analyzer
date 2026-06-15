"""
parser.py - Extracts raw text from various file formats.
Supports: PDF, DOCX, TXT, CSV, JSON, YAML, XLSX, PNG, JPEG/JPG
"""

import io
import csv
import json
import yaml
import traceback

# Image formats supported via OCR
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "bmp", "tiff", "tif", "webp", "gif"}


def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    Given a filename and raw bytes, return the extracted plain text.
    Raises ValueError if the file type is unsupported.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext == "pdf":
            return _parse_pdf(file_bytes)
        elif ext in ("docx", "doc"):
            return _parse_docx(file_bytes)
        elif ext == "txt":
            return file_bytes.decode("utf-8", errors="replace")
        elif ext == "csv":
            return _parse_csv(file_bytes)
        elif ext == "json":
            return _parse_json(file_bytes)
        elif ext in ("yaml", "yml"):
            return _parse_yaml(file_bytes)
        elif ext in ("xlsx", "xls"):
            return _parse_xlsx(file_bytes)
        elif ext in IMAGE_EXTENSIONS:
            return _parse_image(file_bytes, ext)
        elif ext in ("md", "rst", "log", "conf", "ini", "toml"):
            return file_bytes.decode("utf-8", errors="replace")
        elif ext == "":
            # No extension: try raw UTF-8 decode as fallback
            try:
                text = file_bytes.decode("utf-8", errors="replace")
                if len(text.strip()) > 10:
                    return text
            except Exception:
                pass
            raise ValueError(
                "Could not decode extension-less file as UTF-8 text. "
                "Supported types: pdf, docx, doc, txt, csv, json, yaml, yml, "
                "xlsx, xls, md, rst, log, conf, ini, toml, png, jpg, jpeg, bmp, tiff, tif, webp"
            )
        else:
            raise ValueError(
                f"Unsupported file type: .{ext}. "
                "Supported types: pdf, docx, doc, txt, csv, json, yaml, yml, "
                "xlsx, xls, md, rst, log, conf, ini, toml, png, jpg, jpeg, bmp, tiff, tif, webp"
            )
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to parse file '{filename}': {e}\n{traceback.format_exc()}")


def _ocr_page_image(pdf_data: bytes, page_index: int) -> str:
    """
    Render a single PDF page to a PIL Image and run Tesseract OCR on it.
    Applies the same preprocessing as _parse_image (upscale, contrast, sharpness, grayscale).
    Returns the OCR text or empty string if dependencies are unavailable.
    """
    # Try to import a PDF rendering library (fitz preferred, pdf2image as fallback)
    fitz_mod = None
    pdf2image_mod = None
    try:
        import fitz as _fitz
        fitz_mod = _fitz
    except ImportError:
        try:
            from pdf2image import convert_from_bytes as _convert
            pdf2image_mod = _convert
        except ImportError:
            # Neither fitz nor pdf2image available — cannot render page
            return ""

    try:
        import pytesseract
        from PIL import Image, ImageEnhance
    except ImportError:
        return ""

    # Render page to PIL Image
    img = None
    if fitz_mod is not None:
        try:
            doc = fitz_mod.open(stream=pdf_data, filetype="pdf")
            page = doc.load_page(page_index)
            # Render at 2x resolution for better OCR
            mat = fitz_mod.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            doc.close()
        except Exception:
            img = None

    if img is None and pdf2image_mod is not None:
        try:
            # Convert only the specific page (1-indexed for pdf2image)
            images = pdf2image_mod(
                pdf_data,
                first_page=page_index + 1,
                last_page=page_index + 1,
                dpi=200
            )
            if images:
                img = images[0]
        except Exception:
            return ""

    if img is None:
        return ""

    if img is None:
        return ""

    # Preprocessing: same as _parse_image
    # Convert to RGB if needed
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    # Upscale small images
    w, h = img.size
    if w < 1000 or h < 1000:
        scale = max(1000 / w, 1000 / h)
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    # Enhance contrast and sharpness
    img = ImageEnhance.Contrast(img).enhance(1.5)
    img = ImageEnhance.Sharpness(img).enhance(2.0)

    # Convert to grayscale
    img = img.convert("L")

    # Point pytesseract at the known Windows install path if not already on PATH
    import shutil
    if not shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = (
            r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        )

    # Run OCR
    try:
        text = pytesseract.image_to_string(img, config="--psm 6")
        return text.strip()
    except Exception:
        return ""


def _parse_pdf(data: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        # If extracted text is too short (< 10 chars), the page is likely
        # image-only — attempt per-page OCR as fallback.
        if len(text.strip()) < 10:
            ocr_text = _ocr_page_image(data, i)
            if ocr_text:
                text = ocr_text
        if text.strip():
            pages.append(text)
    return "\n".join(pages)


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _parse_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    lines = []
    for row in reader:
        lines.append(", ".join(row))
    return "\n".join(lines)


def _parse_json(data: bytes) -> str:
    obj = json.loads(data.decode("utf-8", errors="replace"))
    return json.dumps(obj, indent=2)


def _parse_yaml(data: bytes) -> str:
    obj = yaml.safe_load(data.decode("utf-8", errors="replace"))
    return yaml.dump(obj, default_flow_style=False)


def _parse_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)


def _parse_image(data: bytes, ext: str) -> str:
    """
    Extract text from an image using Tesseract OCR via pytesseract.
    Applies preprocessing with Pillow to improve OCR accuracy.
    Raises a clear RuntimeError if Tesseract is not installed.
    """
    try:
        import pytesseract
        from PIL import Image, ImageFilter, ImageEnhance
    except ImportError:
        raise RuntimeError(
            "Image OCR requires 'Pillow' and 'pytesseract'. "
            "Run: pip install Pillow pytesseract"
        )

    # Open image
    try:
        img = Image.open(io.BytesIO(data))
    except Exception as e:
        raise RuntimeError(f"Could not open image file: {e}")

    # Convert to RGB if needed (handles RGBA, palette, etc.)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    # Preprocessing: upscale small images, enhance contrast, sharpen
    w, h = img.size
    if w < 1000 or h < 1000:
        scale = max(1000 / w, 1000 / h)
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    # Enhance contrast and sharpness for better OCR
    img = ImageEnhance.Contrast(img).enhance(1.5)
    img = ImageEnhance.Sharpness(img).enhance(2.0)

    # Convert to grayscale for Tesseract
    img = img.convert("L")

    # Point pytesseract at the known Windows install path if not already on PATH
    import shutil, os
    if not shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    # Run OCR
    try:
        text = pytesseract.image_to_string(img, config="--psm 6")
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Tesseract OCR engine is not installed or not on PATH.\n"
            "Install it from: https://github.com/UB-Mannheim/tesseract/wiki\n"
            "Then add the install directory to your system PATH and restart the app."
        )
    except Exception as e:
        raise RuntimeError(f"OCR failed: {e}")

    text = text.strip()
    if not text:
        raise RuntimeError(
            "OCR could not extract any text from the image. "
            "Try a higher-resolution image or ensure the text is clearly visible."
        )

    return text
